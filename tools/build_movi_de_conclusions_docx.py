#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "MOVI_DE_CONCLUSIONS_AND_NEXT_STEPS.docx"
NAVY = "17324D"; BLUE = "2E75B6"; LIGHT = "EAF2F8"; GRAY = "F3F5F7"; AMBER = "FFF1CC"; INK = "17212B"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for side, value in (("top",top),("bottom",bottom),("start",start),("end",end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None: node = OxmlElement(f"w:{side}"); tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tblPr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w/1440); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)
            tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW")); tcW.set(qn("w:w"), str(w)); tcW.set(qn("w:type"), "dxa")

def set_font(run, size=None, bold=None, color=INK, italic=None):
    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), size=10.5); return p

def add_number(doc, text):
    p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(5); p.paragraph_format.line_spacing = 1.1
    set_font(p.add_run(text), size=10.5); return p

doc = Document(); section = doc.sections[0]
section.page_width = Inches(8.5); section.page_height = Inches(11)
section.top_margin = Inches(0.75); section.bottom_margin = Inches(0.7); section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)
section.header_distance = Inches(0.35); section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
for name, size, before, after in (("Heading 1",16,14,7),("Heading 2",13,10,5)):
    st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(BLUE)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

header = section.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("MOVi-D/E CAMERA-POSE EXTENSION  |  FINAL BRIEF"), size=8.5, bold=True, color="637083")
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Synthetic within-video matching study  •  Locked release 2026-08-25"), size=8, color="637083")

p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
set_font(p.add_run("RESEARCH CONCLUSIONS"), size=9, bold=True, color=BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
set_font(p.add_run("Camera pose improves geometry-based matching under camera motion"), size=22, bold=True, color=NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14)
set_font(p.add_run("MOVi-D fixed-camera control • MOVi-E moving-camera confirmation • Transfer and pose-noise evidence"), size=11.5, italic=True, color="52606D")

callout=doc.add_table(rows=1, cols=1); set_table_geometry(callout,[9360]); shade(callout.cell(0,0),LIGHT)
cp=callout.cell(0,0).paragraphs[0]; cp.paragraph_format.space_after=Pt(0)
set_font(cp.add_run("BOTTOM LINE  "), size=10.5, bold=True, color=NAVY)
set_font(cp.add_run("The predeclared primary hypothesis passed: pose-aligned geometry improved AUROC over camera-space geometry on moving-camera MOVi-E. The gain is statistically supported, small in absolute size, and strongest in higher-motion strata."), size=10.5)

doc.add_heading("What the experiment found", level=1)
rows = [
    ("MOVi-E primary", "+0.003057", "0.001131 to 0.005166", "Supported"),
    ("Translation high vs low", "+0.005408", "0.001037 to 0.011271", "Supported secondary"),
    ("Rotation high vs low", "+0.006070", "0.001185 to 0.011617", "Supported secondary"),
    ("MOVi-D fixed camera", "+0.001615", "-0.002037 to 0.004542", "Consistent; not equivalence"),
    ("D→E transfer vs in-domain D", "-0.000154", "-0.000983 to 0.000934", "AUROC transfers"),
]
t=doc.add_table(rows=1,cols=4); headers=["Comparison","AUROC difference","Paired 95% CI","Interpretation"]
for i,h in enumerate(headers):
    shade(t.rows[0].cells[i],BLUE); set_font(t.rows[0].cells[i].paragraphs[0].add_run(h),size=9,bold=True,color="FFFFFF")
for row in rows:
    cells=t.add_row().cells
    for i,v in enumerate(row): set_font(cells[i].paragraphs[0].add_run(v),size=9)
set_table_geometry(t,[2600,1600,2300,2860])

add_bullet(doc,"At the fixed operating point, D reduced MOVi-E false-match rate by 0.002 and increased recall by 0.010 versus C, but both paired confidence intervals included zero; the operating-point claim remains unresolved.")
add_bullet(doc,"Shuffled pose removed the correct-pose advantage, and stronger injected pose noise degraded performance. Together with the motion-stratum result, this supports a camera-coordinate mechanism rather than pose acting as a label shortcut.")
add_bullet(doc,"MOVi-D-to-MOVi-E transfer preserved discrimination, but false-match rate was 0.006 higher than in-domain D (95% CI 0.001005 to 0.011952) at the MOVi-D threshold. Cross-domain threshold calibration matters.")

doc.add_heading("What the result does not establish", level=1)
limit=doc.add_table(rows=1,cols=1); set_table_geometry(limit,[9360]); shade(limit.cell(0,0),AMBER)
lp=limit.cell(0,0).paragraphs[0]; set_font(lp.add_run("ORACLE-GEOMETRY LIMITATION. "),size=10.5,bold=True,color=NAVY)
set_font(lp.add_run("The experiment uses simulator masks, depth, intrinsics, and camera poses. It isolates whether correct pose alignment can help; it does not demonstrate a deployable system using estimated masks, monocular depth, or visual-inertial pose."),size=10.5)
lp=limit.cell(0,0).add_paragraph(); set_font(lp.add_run("FIXED-CAMERA LIMITATION. "),size=10.5,bold=True,color=NAVY)
set_font(lp.add_run("MOVi-D is a synthetic fixed-camera control, not evidence for every static-camera setting. MOVi-D and MOVi-E are not paired counterfactual renders of identical scenes, so the cross-dataset difference-in-differences remains descriptive."),size=10.5)

add_bullet(doc,"Both datasets are synthetic and all benchmark pairs are within-video; no claim is made about cross-video re-identification, real sensor noise, rolling shutter, dynamic backgrounds, or camera-pose estimator failures.")
add_bullet(doc,"Failure cases remain concentrated around occlusion-truncated masks, matched-category distractors, long temporal gaps, camera-space misalignment, and visible-surface/depth instability.")

doc.add_heading("Recommended next steps", level=1)
for text in [
    "Replace oracle pose, depth, and masks one component at a time with estimated inputs and attribute the resulting accuracy loss.",
    "Add correlated trajectory noise, drift, and synchronization error; the current noise grid is a controlled sensitivity study, not a full estimator model.",
    "Render matched fixed- and moving-camera versions of identical scenes to create a true camera-motion counterfactual.",
    "Test target-threshold recalibration and unsupervised adaptation while keeping the MOVi-E test pool locked.",
    "Validate on real handheld and fixed-camera footage before adding cross-video matching or deployment claims.",
]: add_number(doc,text)

doc.add_heading("Release pointers", level=1)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)
set_font(p.add_run("Full tables: "),size=9.5,bold=True); set_font(p.add_run("docs/MOVI_DE_RESULTS_TABLES.xlsx  •  "),size=9.5)
set_font(p.add_run("Failure gallery: "),size=9.5,bold=True); set_font(p.add_run("failure_gallery/movi_de_phase10/failure_gallery.html  •  "),size=9.5)
set_font(p.add_run("Machine-readable evidence: "),size=9.5,bold=True); set_font(p.add_run("results/movi_de_phase9/phase9_criteria_evaluation.json"),size=9.5)

doc.core_properties.title = "MOVi-D/E Camera-Pose Extension: Conclusions and Next Steps"
doc.core_properties.subject = "Final research brief"
doc.core_properties.author = "Research team"
doc.save(OUT)
print(OUT)
